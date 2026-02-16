"""
DOC/DOCX Loader for Microsoft Word document processing.

This loader extracts text and images from Word documents (.doc, .docx).
It supports:
- Text extraction with paragraph and table preservation
- Image extraction from the document
- Configurable chunking of plain text
- Both legacy .doc and modern .docx formats

The loader stores extracted images in a scratch folder within the data directory.
"""

from __future__ import annotations

import mimetypes
import os
import re
import subprocess
import tempfile
import time
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from rakam_systems_tools.utils import logging
from rakam_systems_core.interfaces.loader import Loader
from rakam_systems_vectorstore.components.chunker import AdvancedChunker
from rakam_systems_vectorstore.core import Node, NodeMetadata, VSFile

logger = logging.getLogger(__name__)


class DocLoader(Loader):
    """
    Word document loader for .doc and .docx files.

    This loader provides Word document processing with support for:
    - Text extraction with paragraph and table preservation
    - Image extraction from document archive (DOCX only)
    - Advanced text chunking
    - Both legacy .doc and modern .docx formats

    For .docx files, images are extracted and saved to a scratch directory.
    For legacy .doc files, text extraction is attempted via python-docx or
    falls back to antiword/textutil if available.
    """

    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.doc', '.docx', '.DOC', '.DOCX'}

    # MIME types for Word documents
    MIME_TYPES = {
        'application/msword',  # .doc
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
    }

    # Default configuration
    DEFAULT_EMBED_MODEL_ID = "sentence-transformers/all-MiniLM-L6-v2"
    DEFAULT_CHUNK_SIZE = 2048
    DEFAULT_CHUNK_OVERLAP = 128
    DEFAULT_IMAGE_PATH = "data/ingestion_image/"  # Default path for extracted images

    def __init__(
        self,
        name: str = "doc_loader",
        config: Optional[Dict[str, Any]] = None
    ):
        """
        Initialize DOC/DOCX loader.

        Args:
            name: Component name
            config: Optional configuration with keys:
                - embed_model_id: HuggingFace model ID for tokenization (default: "sentence-transformers/all-MiniLM-L6-v2")
                - chunk_size: Maximum tokens per chunk (default: 2048)
                - chunk_overlap: Overlap between chunks in tokens (default: 128)
                - min_sentences_per_chunk: Minimum sentences per chunk (default: 1)
                - tokenizer: Tokenizer for chunking (default: "character")
                - save_images: Whether to save images to disk (default: True)
                - image_path: Path to save images (default: None, uses INGESTION_IMAGE_PATH env var or "data/ingestion_image/")
                - scratch_folder_name: Name of scratch folder (default: "scratch")
                - include_images_in_text: Whether to add image references to text (default: True)
                - extract_tables: Whether to extract table content (default: True)
                - preserve_formatting: Whether to preserve basic formatting markers (default: False)
        """
        super().__init__(name=name, config=config)

        # Extract configuration
        config = config or {}
        self._save_images = config.get('save_images', True)
        self._image_path = config.get('image_path') or os.getenv(
            'INGESTION_IMAGE_PATH', self.DEFAULT_IMAGE_PATH)
        self._scratch_folder_name = config.get(
            'scratch_folder_name', 'scratch')
        self._include_images_in_text = config.get(
            'include_images_in_text', True)
        self._extract_tables = config.get('extract_tables', True)
        self._preserve_formatting = config.get('preserve_formatting', False)

        # Chunking configuration
        self._chunk_size = config.get('chunk_size', self.DEFAULT_CHUNK_SIZE)
        self._chunk_overlap = config.get(
            'chunk_overlap', self.DEFAULT_CHUNK_OVERLAP)
        self._min_sentences_per_chunk = config.get(
            'min_sentences_per_chunk', 1)
        self._tokenizer = config.get('tokenizer', 'character')

        # Initialize advanced chunker
        embed_model_id = config.get(
            'embed_model_id', self.DEFAULT_EMBED_MODEL_ID)
        self._chunker = AdvancedChunker(
            embed_model_id=embed_model_id,
            strategy="default"
        )

        # Store last extraction info for image tracking
        self._last_scratch_dir = None
        self._last_image_files = []
        self._image_path_mapping: Dict[str, str] = {}

        logger.info(
            f"Initialized DocLoader with chunk_size={self._chunk_size}, chunk_overlap={self._chunk_overlap}, image_path={self._image_path}")

    def run(self, source: str) -> List[str]:
        """
        Execute the primary operation for the component.

        This method satisfies the BaseComponent abstract method requirement
        and delegates to load_as_chunks.

        Args:
            source: Path to DOC/DOCX file

        Returns:
            List of text chunks extracted from the document
        """
        return self.load_as_chunks(source)

    def load_as_text(
        self,
        source: Union[str, Path],
    ) -> str:
        """
        Load Word document and return as a single text string.

        This method extracts all text from the document and returns it as a single
        string without chunking. Useful when you need the full document text.

        Args:
            source: Path to DOC/DOCX file

        Returns:
            Full text content of the document as a single string

        Raises:
            FileNotFoundError: If source file doesn't exist
            ValueError: If source is not a Word document
            Exception: If document processing fails
        """
        # Convert Path to string
        if isinstance(source, Path):
            source = str(source)

        # Validate file exists
        if not os.path.isfile(source):
            raise FileNotFoundError(f"File not found: {source}")

        # Validate file is a Word document
        source_path = Path(source)
        if not self._is_doc_file(source):
            raise ValueError(
                f"File is not a Word document: {source} (extension: {source_path.suffix})")

        logger.info(
            f"Loading Word document as text: {source_path.name} (extension: {source_path.suffix})")
        start_time = time.time()

        try:
            # Create scratch directory in data folder
            scratch_dir = self._get_scratch_dir(source)
            self._last_scratch_dir = scratch_dir

            # Extract images if enabled (DOCX only)
            image_files = []
            if self._save_images and self._is_docx_file(source):
                image_dir = self._get_image_path(source)
                image_files = self._extract_images(source, Path(image_dir))
                self._last_image_files = image_files
                logger.info(
                    f"Extracted {len(image_files)} images from document")

            # Extract text from document
            if self._is_docx_file(source):
                full_text = self._extract_text_docx(source)
            else:
                full_text = self._extract_text_doc(source)

            # Add image references if enabled
            if self._include_images_in_text and image_files:
                full_text = self._add_image_references_to_text(
                    full_text, image_files)

            elapsed = time.time() - start_time
            logger.info(
                f"Document loaded as text in {elapsed:.2f}s: {len(full_text)} characters")

            return full_text

        except Exception as e:
            logger.error(f"Error loading document as text {source}: {e}")
            raise

    def load_as_chunks(
        self,
        source: Union[str, Path],
    ) -> List[str]:
        """
        Load Word document and return as a list of text chunks.

        This method extracts text from the document, processes it with the configured
        chunker, and returns a list of text chunks. Each chunk optionally includes
        image references.

        Args:
            source: Path to DOC/DOCX file

        Returns:
            List of text chunks extracted from the document

        Raises:
            FileNotFoundError: If source file doesn't exist
            ValueError: If source is not a Word document
            Exception: If document processing fails
        """
        # Convert Path to string
        if isinstance(source, Path):
            source = str(source)

        # Validate file exists
        if not os.path.isfile(source):
            raise FileNotFoundError(f"File not found: {source}")

        # Validate file is a Word document
        source_path = Path(source)
        if not self._is_doc_file(source):
            raise ValueError(
                f"File is not a Word document: {source} (extension: {source_path.suffix})")

        logger.info(
            f"Loading Word document file: {source_path.name} (extension: {source_path.suffix})")
        start_time = time.time()

        try:
            # Get full text
            full_text = self.load_as_text(source)

            # Chunk the text using AdvancedChunker's chunk_text method
            text_chunks = self._chunk_text(full_text)

            elapsed = time.time() - start_time
            logger.info(
                f"Document processed in {elapsed:.2f}s: {len(text_chunks)} chunks")

            return text_chunks

        except Exception as e:
            logger.error(f"Error processing document {source}: {e}")
            raise

    def load_as_nodes(
        self,
        source: Union[str, Path],
        source_id: Optional[str] = None,
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> List[Node]:
        """
        Load Word document and return as Node objects with metadata.

        Args:
            source: Path to DOC/DOCX file
            source_id: Optional source identifier (defaults to file path)
            custom_metadata: Optional custom metadata to attach to nodes

        Returns:
            List of Node objects with text chunks and metadata
        """
        # Convert Path to string
        if isinstance(source, Path):
            source = str(source)

        # Load text chunks
        chunks = self.load_as_chunks(source)

        # Determine source ID
        if source_id is None:
            source_id = source

        # Create nodes with metadata
        nodes = []
        for idx, chunk in enumerate(chunks):
            metadata = NodeMetadata(
                source_file_uuid=source_id,
                position=idx,
                custom=custom_metadata or {}
            )
            node = Node(content=chunk, metadata=metadata)
            nodes.append(node)

        logger.info(f"Created {len(nodes)} nodes from document: {source}")
        return nodes

    def load_as_vsfile(
        self,
        file_path: Union[str, Path],
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> VSFile:
        """
        Load Word document and return as VSFile object.

        Args:
            file_path: Path to DOC/DOCX file
            custom_metadata: Optional custom metadata

        Returns:
            VSFile object with nodes

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file is not a Word document
        """
        if isinstance(file_path, Path):
            file_path = str(file_path)

        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        if not self._is_doc_file(file_path):
            raise ValueError(f"File is not a Word document: {file_path}")

        # Create VSFile
        vsfile = VSFile(file_path)

        # Load and create nodes
        nodes = self.load_as_nodes(
            file_path, str(vsfile.uuid), custom_metadata)
        vsfile.nodes = nodes
        vsfile.processed = True

        logger.info(
            f"Created VSFile with {len(nodes)} nodes from: {file_path}")
        return vsfile

    def _is_doc_file(self, file_path: str) -> bool:
        """
        Check if file is a Word document based on extension and magic bytes.

        Args:
            file_path: Path to file

        Returns:
            True if file is a Word document, False otherwise
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        # First check extension
        if suffix not in {'.doc', '.docx'}:
            logger.debug(
                f"File {path.name} rejected: extension '{suffix}' is not .doc or .docx")
            return False

        # Additional safety check: verify it's not a PDF by checking magic bytes
        try:
            with open(file_path, 'rb') as f:
                magic_bytes = f.read(4)
                # PDF files start with %PDF (0x25504446)
                if magic_bytes.startswith(b'%PDF'):
                    logger.error(
                        f"File {path.name} has .doc/.docx extension but is actually a PDF!")
                    return False
        except Exception as e:
            logger.warning(f"Could not read magic bytes from {path.name}: {e}")

        return True

    def _is_docx_file(self, file_path: str) -> bool:
        """
        Check if file is specifically a .docx file.

        Args:
            file_path: Path to file

        Returns:
            True if file is a .docx, False otherwise
        """
        path = Path(file_path)
        return path.suffix.lower() == '.docx'

    def _get_scratch_dir(self, source_path: str) -> Path:
        """
        Get scratch directory for storing extracted files.

        The scratch directory is created inside the data folder relative to the source file.

        Args:
            source_path: Path to source document file

        Returns:
            Path to scratch directory
        """
        source = Path(source_path)

        # Find data folder - assume it's a parent of the source or sibling
        if 'data' in source.parts:
            # Navigate to data folder
            data_folder = source
            while data_folder.name != 'data' and data_folder.parent != data_folder:
                data_folder = data_folder.parent
        else:
            # Use parent directory and create/use data folder
            data_folder = source.parent / 'data'

        # Create scratch directory inside data folder
        scratch_dir = data_folder / self._scratch_folder_name
        scratch_dir.mkdir(parents=True, exist_ok=True)

        logger.debug(f"Using scratch directory: {scratch_dir}")
        return scratch_dir

    def _get_image_path(self, source_path: str) -> str:
        """
        Get the path where images should be extracted.

        Uses the configured image path (from config, env var, or default).
        Creates a subdirectory based on the source document filename.

        Args:
            source_path: Path to source document file

        Returns:
            Absolute path to image extraction directory
        """
        source = Path(source_path)
        doc_filename = source.stem

        # Create base image path
        base_path = Path(self._image_path)

        # Create subdirectory for this document
        image_dir = base_path / doc_filename
        image_dir.mkdir(parents=True, exist_ok=True)

        logger.debug(f"Using image path: {image_dir}")
        return str(image_dir)

    def get_image_path_mapping(self) -> Dict[str, str]:
        """
        Get the mapping of image paths.

        Returns:
            Dictionary mapping image filenames to absolute paths on disk
        """
        return self._image_path_mapping.copy()

    def get_image_absolute_path(self, image_filename: str) -> Optional[str]:
        """
        Get the absolute file path for an image.

        Args:
            image_filename: The image filename

        Returns:
            Absolute path to the image file, or None if not found
        """
        return self._image_path_mapping.get(image_filename)

    def _extract_text_docx(self, docx_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.

        Args:
            docx_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        # Additional safety check: ensure this is actually a DOCX/DOC file
        file_path = Path(docx_path)
        if file_path.suffix.lower() not in ['.doc', '.docx']:
            raise ValueError(
                f"File is not a Word document: {docx_path} (extension: {file_path.suffix})")

        try:
            from docx import Document
        except ImportError:
            logger.error(
                "python-docx is required for DOCX support. Install with: pip install python-docx")
            raise ImportError("python-docx is required for DOCX support")

        try:
            doc = Document(docx_path)

            # Check if document was successfully loaded
            if doc is None:
                raise ValueError(f"Failed to load document: {docx_path}")

            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text
                if para_text.strip():
                    # Optionally preserve formatting markers
                    if self._preserve_formatting:
                        # Add heading markers based on style
                        style_name = paragraph.style.name if paragraph.style else ""
                        if style_name.startswith("Heading"):
                            level = style_name[-1] if style_name[-1].isdigit() else "1"
                            para_text = f"{'#' * int(level)} {para_text}"
                    text_parts.append(para_text)

            # Extract tables if enabled
            if self._extract_tables and hasattr(doc, 'tables') and doc.tables is not None:
                for table in doc.tables:
                    table_text = self._extract_table_text(table)
                    if table_text.strip():
                        text_parts.append(table_text)

            full_text = '\n\n'.join(text_parts)
            logger.debug(f"Extracted {len(full_text)} characters from DOCX")
            return full_text

        except Exception as e:
            # Check if this might be a PDF file mistakenly routed here
            if file_path.suffix.lower() == '.pdf' or 'pdf' in str(e).lower():
                raise ValueError(
                    f"File appears to be a PDF, not a Word document: {docx_path}. Error: {e}")
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise

    def _extract_table_text(self, table) -> str:
        """
        Extract text from a Word table.

        Args:
            table: python-docx Table object

        Returns:
            Table content as formatted text
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(' | '.join(cells))
        return '\n'.join(rows)

    def _extract_text_doc(self, doc_path: str) -> str:
        """
        Extract text from legacy .doc file.

        This method tries multiple approaches:
        1. Use python-docx (may work for some .doc files)
        2. Use antiword if available (Linux/macOS)
        3. Use textutil if available (macOS)

        Args:
            doc_path: Path to .doc file

        Returns:
            Extracted text content
        """
        # Safety check: verify this is not a PDF file
        try:
            with open(doc_path, 'rb') as f:
                magic_bytes = f.read(4)
                if magic_bytes.startswith(b'%PDF'):
                    raise ValueError(
                        f"File {doc_path} is a PDF, not a Word document. It should not be processed by DocLoader.")
        except IOError:
            pass  # If we can't read the file, let the extraction methods handle it

        # First, try python-docx (works for some .doc files that are actually .docx in disguise)
        try:
            return self._extract_text_docx(doc_path)
        except Exception as e:
            logger.debug(f"python-docx failed for .doc file: {e}")

        # Try antiword (available on Linux and some macOS systems)
        try:
            result = subprocess.run(
                ['antiword', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                logger.debug("Successfully extracted text using antiword")
                return result.stdout
        except FileNotFoundError:
            logger.debug("antiword not available")
        except subprocess.TimeoutExpired:
            logger.warning("antiword timed out")
        except Exception as e:
            logger.debug(f"antiword failed: {e}")

        # Try textutil (macOS)
        try:
            with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
                tmp_path = tmp.name

            result = subprocess.run(
                ['textutil', '-convert', 'txt', '-output', tmp_path, doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                with open(tmp_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                os.unlink(tmp_path)
                logger.debug("Successfully extracted text using textutil")
                return content
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except FileNotFoundError:
            logger.debug("textutil not available")
        except subprocess.TimeoutExpired:
            logger.warning("textutil timed out")
        except Exception as e:
            logger.debug(f"textutil failed: {e}")

        # If all methods fail, raise an error
        raise RuntimeError(
            f"Could not extract text from .doc file: {doc_path}. "
            "Install 'antiword' (Linux/macOS) or use macOS with 'textutil', "
            "or convert the file to .docx format."
        )

    def _extract_images(self, docx_path: str, output_dir: Path) -> List[str]:
        """
        Extract all images from a DOCX file.

        DOCX files are ZIP archives with images stored in the word/media/ directory.

        Args:
            docx_path: Path to the DOCX file
            output_dir: Directory to save extracted images

        Returns:
            List of paths to extracted image files
        """
        docx_path = Path(docx_path)
        extracted_files = []

        # Clear previous mapping
        self._image_path_mapping.clear()

        try:
            # DOCX files are ZIP archives
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # Images are stored in word/media/ directory
                for file_info in zip_ref.filelist:
                    # Extract only image files from media folder
                    if file_info.filename.startswith('word/media/') and not file_info.is_dir():
                        # Get the filename
                        filename = Path(file_info.filename).name

                        # Check if it's an image based on extension
                        img_extensions = {'.png', '.jpg', '.jpeg',
                                          '.gif', '.bmp', '.tiff', '.emf', '.wmf'}
                        if Path(filename).suffix.lower() in img_extensions:
                            # Extract the file
                            extracted_path = output_dir / filename
                            with zip_ref.open(file_info) as source, open(extracted_path, 'wb') as target:
                                target.write(source.read())
                            extracted_files.append(str(extracted_path))

                            # Build image path mapping
                            self._image_path_mapping[filename] = str(
                                extracted_path)
                            logger.debug(f"Extracted image: {extracted_path}")

            logger.info(f"Extracted {len(extracted_files)} images from DOCX")
            logger.info(
                f"Built image path mapping with {len(self._image_path_mapping)} images")

        except Exception as e:
            logger.warning(f"Failed to extract images from DOCX: {e}")

        return extracted_files

    def _add_image_references_to_text(self, text: str, image_files: List[str]) -> str:
        """
        Add image references to the extracted text.

        Args:
            text: Extracted text content
            image_files: List of extracted image file paths

        Returns:
            Text with appended image references
        """
        if not image_files:
            return text

        # Add image references at the end of the text
        image_refs = "\n\n--- Embedded Images ---\n"
        for img_path in image_files:
            img_name = Path(img_path).name
            image_refs += f"\n![{img_name}]({img_path})"

        return text + image_refs

    def _chunk_text(self, text: str) -> List[str]:
        """
        Chunk text using AdvancedChunker's chunk_text method.

        This method uses chunk_text() which is specifically designed for plain text strings.

        Args:
            text: Full text to chunk

        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []

        try:
            # Use AdvancedChunker's chunk_text method for plain text
            chunk_dicts = self._chunker.chunk_text(
                text=text,
                chunk_size=self._chunk_size,
                chunk_overlap=self._chunk_overlap,
                min_sentences_per_chunk=self._min_sentences_per_chunk,
                tokenizer=self._tokenizer
            )

            # Extract just the text from the chunk dictionaries
            text_chunks = [chunk_dict['text'] for chunk_dict in chunk_dicts]

            logger.info(f"Chunked text into {len(text_chunks)} chunks")
            return text_chunks

        except Exception as e:
            logger.warning(f"Failed to chunk text with AdvancedChunker: {e}")
            # Fall back to returning the whole text as a single chunk
            logger.info("Falling back to single chunk")
            return [text]


def create_doc_loader(
    chunk_size: int = 2048,
    chunk_overlap: int = 128,
    min_sentences_per_chunk: int = 1,
    tokenizer: str = "character",
    embed_model_id: str = "sentence-transformers/all-MiniLM-L6-v2",
    save_images: bool = True,
    scratch_folder_name: str = 'scratch',
    include_images_in_text: bool = True,
    extract_tables: bool = True,
    preserve_formatting: bool = False
) -> DocLoader:
    """
    Factory function to create a Word document loader.

    Args:
        chunk_size: Maximum tokens per chunk (default: 2048)
        chunk_overlap: Overlap between chunks in tokens (default: 128)
        min_sentences_per_chunk: Minimum sentences per chunk (default: 1)
        tokenizer: Tokenizer for chunking - "character", "gpt2", or HuggingFace model (default: "character")
        embed_model_id: HuggingFace model ID for tokenization (default: "sentence-transformers/all-MiniLM-L6-v2")
        save_images: Whether to save extracted images (default: True)
        scratch_folder_name: Name of scratch folder in data directory (default: "scratch")
        include_images_in_text: Whether to include image references in text (default: True)
        extract_tables: Whether to extract table content (default: True)
        preserve_formatting: Whether to preserve basic formatting markers (default: False)

    Returns:
        Configured DOC/DOCX loader

    Example:
        >>> loader = create_doc_loader(chunk_size=1024, chunk_overlap=64)
        >>> chunks = loader.run("data/document.docx")
        >>> print(f"Extracted {len(chunks)} chunks")

        >>> # Create loader without image references
        >>> loader = create_doc_loader(include_images_in_text=False)
        >>> chunks = loader.run("data/document.docx")

        >>> # Load as nodes for vector store
        >>> loader = create_doc_loader()
        >>> nodes = loader.load_as_nodes("data/report.docx", custom_metadata={"category": "reports"})
    """
    config = {
        'chunk_size': chunk_size,
        'chunk_overlap': chunk_overlap,
        'min_sentences_per_chunk': min_sentences_per_chunk,
        'tokenizer': tokenizer,
        'embed_model_id': embed_model_id,
        'save_images': save_images,
        'scratch_folder_name': scratch_folder_name,
        'include_images_in_text': include_images_in_text,
        'extract_tables': extract_tables,
        'preserve_formatting': preserve_formatting
    }

    return DocLoader(config=config)


__all__ = ["DocLoader", "create_doc_loader"]
